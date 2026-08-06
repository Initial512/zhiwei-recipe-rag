from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("知味AI_项目经理_STAR项目经历.docx")
FONT = "Microsoft YaHei"
NAVY = RGBColor(31, 78, 121)
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(89, 89, 89)


def set_font(run, size=10.5, bold=False, color=BLACK, underline=False):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = color
    run.underline = underline


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_margins(cell, top=90, start=100, bottom=90, end=100):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color="B7C9D6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def paragraph(doc, before=0, after=4, line=1.35, left=0, first=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    pf.left_indent = Cm(left)
    pf.first_line_indent = Cm(first)
    return p


def add_label_paragraph(doc, label, text, after=4):
    p = paragraph(doc, after=after)
    set_font(p.add_run(label), size=10.5, bold=True, color=NAVY)
    set_font(p.add_run(text), size=10.5)
    return p


def add_bullet(doc, label, text):
    p = paragraph(doc, after=3, left=0.48, first=-0.48)
    set_font(p.add_run("■  "), size=10.5, bold=True, color=BLACK)
    set_font(p.add_run(label), size=10.5, bold=True)
    set_font(p.add_run(text), size=10.5)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Cm(1.5)
section.bottom_margin = Cm(1.5)
section.left_margin = Cm(1.7)
section.right_margin = Cm(1.7)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)

# Header-style project line, following the supplied resume example.
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [Cm(4.2), Cm(8.4), Cm(3.1)]
for cell, width in zip(table.rows[0].cells, widths):
    cell.width = width
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(cell, top=75, start=100, bottom=75, end=100)
    set_cell_border(cell)
    shade(cell, "F6F9FC")

first, link, date = table.rows[0].cells
p = first.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(p.add_run("知味 AI（项目经理）"), size=12, bold=True, color=NAVY)
p = link.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run("https://github.com/Initial512/zhiwei-recipe-rag"), size=9.5, color=NAVY, underline=True)
p = date.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(p.add_run("2026.07 - 至今"), size=10, bold=True)

add_label_paragraph(
    doc,
    "技术栈：",
    "Python、FastAPI、React 19、Vite 6、LangChain、OpenAI 兼容 API、Neo4j 5、Milvus 2、BAAI/bge-small-zh-v1.5、Docker Compose、Nginx、SSE"
)

add_label_paragraph(
    doc,
    "项目简介：",
    "面向中文菜谱场景构建智能美食推荐与问答平台。围绕菜谱浏览、分类检索、自然语言推荐和流式 AI 问答等需求，将 Markdown 菜谱资料结构化为 Neo4j 知识图谱，并结合 Milvus 语义检索与大模型生成，提供可解释、可扩展的 GraphRAG 服务闭环。"
)

p = paragraph(doc, before=3, after=4)
set_font(p.add_run("技术亮点："), size=10.5, bold=True, color=NAVY)

add_bullet(
    doc,
    "GraphRAG 检索架构升级：",
    "针对传统“Markdown 文档 + 单一路径向量检索”难以表达菜谱、食材、烹饪步骤和分类关联的问题，设计 Neo4j 知识图谱 + Milvus 语义向量库的 GraphRAG 架构；以图谱承载结构化关系、向量库承载语义召回，为复杂饮食需求提供可解释的检索依据。"
)
add_bullet(
    doc,
    "分层召回与混合检索：",
    "构建“精确菜名命中 → 模糊菜名匹配 → 图谱/向量候选兜底”的三级召回链路；关系型短查询优先进入 Neo4j 图谱检索，泛化需求优先进入 Milvus 语义检索，较长的关系组合问题同时汇聚两类候选，兼顾高频查询的确定性和复杂问题的召回覆盖。"
)
add_bullet(
    doc,
    "结构化解析与融合排序：",
    "将口味、食材、菜品类型等自然语言条件解析为结构化标签，并引入“元数据匹配度 60% + 关键词命中 25% + 向量相关度 15%”的融合排序；支持多条件过滤、食材别名匹配与代表性菜品优先排序，提升推荐结果与用户需求的贴合度。"
)
add_bullet(
    doc,
    "前端查询意图动态分流：",
    "封装并测试 React 端搜索路由模块，通过“怎么做、食材、步骤、菜谱”等疑问标记识别菜谱知识问答；非菜谱问句根据本地候选结果决策结果卡片模式或菜谱知识模式，并以 URL 参数承载页面模式状态，为“直接找菜”和“询问做法”提供差异化交互路径。"
)
add_bullet(
    doc,
    "流式 RAG 问答闭环：",
    "基于 FastAPI 建设检索、详情、推荐和问答接口，通过 SSE 依次推送菜谱来源、生成增量与完成事件；前端同步展示流式答案和来源菜谱卡片，形成“检索召回—大模型生成—结果可追溯”的问答闭环。"
)
add_bullet(
    doc,
    "数据资产与工程化交付：",
    "完成 322 份 Markdown 菜谱、322 张 WebP 图片及图谱实体关系的数据治理；通过 Docker Compose 编排 Neo4j、Milvus、FastAPI 与 Nginx，并配套健康检查、后端单测、前端 lint/build 和配置校验，保障 RAG 服务可部署、可验证。"
)

# Compact footer note with factual boundary.
p = paragraph(doc, before=4, after=0)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(p.add_run("数据与技术描述基于当前代码库扫描整理"), size=8.5, color=GRAY)

doc.core_properties.title = "知味 AI - 项目经理 STAR 项目经历"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT.resolve())
