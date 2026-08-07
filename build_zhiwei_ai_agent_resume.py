from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUT = Path("知味AI_AI_Agent开发_项目经历.docx")
FONT = "微软雅黑"
EN_FONT = "Times New Roman"


def set_font(run, size=10.5, bold=False, font=FONT):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run.font.size = Pt(size)
    run.bold = bold


def add_numbering(doc):
    """Create a real Word bullet numbering definition using the resume's square marker."""
    numbering = doc.part.numbering_part.element
    abstract_id = max(
        [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
        or [-1]
    ) + 1
    num_id = max(
        [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))] or [0]
    ) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    fmt = OxmlElement("w:numFmt")
    fmt.set(qn("w:val"), "bullet")
    level.append(fmt)
    text = OxmlElement("w:lvlText")
    text.set(qn("w:val"), "■")
    level.append(text)
    justification = OxmlElement("w:lvlJc")
    justification.set(qn("w:val"), "left")
    level.append(justification)
    ppr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "480")
    ind.set(qn("w:hanging"), "480")
    ppr.append(ind)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def set_bullet(paragraph, num_id):
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    ppr.append(num_pr)


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def set_table_borders_none(table):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def add_paragraph(doc, before=0, after=4, line=1.0):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    return paragraph


def add_labeled_paragraph(doc, label, text):
    paragraph = add_paragraph(doc, after=4)
    set_font(paragraph.add_run(label), bold=True)
    set_font(paragraph.add_run(text))


def add_highlight(doc, num_id, label, text):
    paragraph = add_paragraph(doc, after=3, line=1.12)
    set_bullet(paragraph, num_id)
    set_font(paragraph.add_run(label), bold=True)
    set_font(paragraph.add_run(text))


doc = Document()
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(0.9)
section.bottom_margin = Cm(0.9)
section.left_margin = Cm(1.25)
section.right_margin = Cm(1.25)

normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
normal.font.size = Pt(10.5)

# Same three-column project heading pattern as the supplied resume.
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
set_table_borders_none(table)
left, right = table.rows[0].cells
left.width = Cm(14.8)
right.width = Cm(4.0)
for cell in (left, right):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)

title = left.paragraphs[0]
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
set_font(title.add_run("知味 AI 智能饮食知识检索与推荐系统（AI Agent 开发）"), size=11.5, bold=True)
date = right.paragraphs[0]
date.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(date.add_run("2026.06 - 至今"), size=10.5)

add_labeled_paragraph(
    doc,
    "技术栈：",
    "Python、FastAPI、LangChain、OpenAI 兼容 API、Neo4j、Milvus、BM25、jieba、rank-bm25、BAAI/bge-small-zh-v1.5、React、Docker Compose、Nginx、SSE",
)
add_labeled_paragraph(
    doc,
    "项目简介：",
    "面向家庭烹饪场景构建“浏览—搜索—推荐—追问”闭环 AI 应用。针对 322 份中文 Markdown 菜谱中食材、分类与烹饪步骤关联难以由单一路径检索覆盖的问题，构建 Neo4j + Milvus 的 GraphRAG 检索 Agent；通过本地菜名直查、三路混合召回、RRF 融合重排和流式生成，提供可解释、可追溯的菜谱检索与饮食建议。",
)

heading = add_paragraph(doc, before=2, after=4)
set_font(heading.add_run("技术亮点："), bold=True)
numbering_id = add_numbering(doc)

add_highlight(
    doc,
    numbering_id,
    "GraphRAG 检索 Agent：",
    "针对传统“文档分块 + 单路向量召回”难以表达菜谱、食材、步骤与分类关联的问题，将结构化关系沉淀至 Neo4j、语义相似度交由 Milvus；复杂问题按查询复杂度规划单跳、多跳或组合子图检索，并按图结构相关度排序，降低关联信息遗漏风险。",
)
add_highlight(
    doc,
    numbering_id,
    "图数据建模与结构化索引：",
    "将 Markdown 菜谱转换为 Neo4j 节点与关系资产，通过 Cypher 导入菜谱、食材、烹饪步骤、难度与分类等实体关系；从图数据构建结构化菜谱文档并分块，携带菜名、节点类型、分类、难度等元数据写入 Milvus，支撑语义召回与条件过滤。",
)
add_highlight(
    doc,
    numbering_id,
    "三路混合检索 + RRF 重排序：",
    "引入 jieba 中文分词与 BM25Okapi，保留单字食材等关键 Token；并行执行双层图谱、Milvus 向量和 BM25 词法召回，以 RRF（k=60）汇聚各路排名，并按菜谱节点去重、记录来源与排名元数据，提升多路候选的排序稳健性。",
)
add_highlight(
    doc,
    numbering_id,
    "动态路由与容错降级：",
    "为高频“找菜”需求提供本地菜名关键词匹配接口，不调用 LLM 即返回候选；通用 RAG 查询由 LLM 分析复杂度、关系密集度、推理需求和实体数，动态选择传统混合检索、GraphRAG 或组合检索；模型分析异常时切换关键词规则，复杂检索无结果或失败时降级至混合检索。",
)
add_highlight(
    doc,
    numbering_id,
    "流式问答与可追溯交互：",
    "基于 FastAPI 建设检索、推荐、详情与问答接口，通过 SSE 按“来源 → 生成增量 → 完成”事件协议推送结果；React 端同步呈现回答与关联菜谱卡片，形成“检索召回—大模型生成—来源回溯”的用户闭环。",
)
add_highlight(
    doc,
    numbering_id,
    "工程化交付与质量保障：",
    "使用 Docker Compose 编排 Neo4j、Milvus、FastAPI 与 Nginx，结合健康检查、10 次/分钟接口限流、后端单测、前端 lint/build 与配置校验，保障多服务 RAG Agent 的可部署、可验证与稳定运行。",
)

doc.core_properties.title = "知味 AI - AI Agent 开发项目经历"
doc.core_properties.author = "Codex"
doc.save(OUT)
print(OUT.resolve())
